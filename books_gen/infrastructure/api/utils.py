from typing import Dict
import os
from books_gen.models.book_models import Book, FormatDownload
from books_gen.config import settings
# Para PDF: reportlab es completamente Python
import markdown
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def convert_json_to_markdown(book: Book) -> str:
    """Convierte el objeto Book a contenido markdown."""
    markdown_content = f"# {book.title}\n\n{book.synopsis}\n\n"
    for chapter in book.index["chapters"]:
        markdown_content += f"## {chapter['title']}\n\n{chapter['description']}\n\n"
        if chapter.get("content"):
            markdown_content += f"{chapter['content']}\n\n"
    
    # Guardar el contenido en un archivo temporal
    temp_file_path = os.path.join(settings.BOOKS_DIR, f"{book.id}.md")
    with open(temp_file_path, "w", encoding="utf-8") as f:
        f.write(markdown_content)
    
    return markdown_content

def convert_markdown_to_html(markdown_content: str, book: Book) -> str:
    """Convierte contenido markdown a HTML."""
    html_content = markdown.markdown(markdown_content, extensions=['fenced_code', 'tables'])
    # Crear HTML con estilo
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{book.title}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 3cm; }}
            h1 {{ font-size: 24pt; color: #333; }}
            h2 {{ font-size: 18pt; color: #444; margin-top: 2em; }}
            p {{ font-size: 12pt; line-height: 1.5; }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    return styled_html

def convert_markdown_to_pdf_with_reportlab(markdown_content: str, book: Book) -> str:
    """
    Convierte contenido markdown a PDF usando reportlab (sin dependencias externas).
    
    Args:
        markdown_content: Contenido en formato markdown.
        book: Objeto Book con los datos del libro.
        
    Returns:
        Ruta del archivo PDF generado.
    """
    output_file_path = os.path.join(settings.BOOKS_DIR, f"{book.id}.pdf")
    
    # Crear el documento
    doc = SimpleDocTemplate(
        output_file_path,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    # Estilos - en lugar de añadir, modificamos los existentes o creamos nuevos con nombres únicos
    styles = getSampleStyleSheet()
    
    # Modificar el estilo Title existente en lugar de añadir uno nuevo
    styles['Title'].fontSize = 24
    styles['Title'].alignment = 1  # Centrado
    styles['Title'].spaceAfter = 12
    
    # Modificar el estilo Heading2 existente
    styles['Heading2'].fontSize = 18
    styles['Heading2'].spaceAfter = 10
    styles['Heading2'].spaceBefore = 20
    
    # Modificar el estilo Normal existente
    styles['Normal'].fontSize = 12
    styles['Normal'].spaceAfter = 6
    
    # También podemos crear estilos adicionales con nombres únicos si es necesario
    styles.add(ParagraphStyle(
        name='BookListItem',  # Nombre único para elementos de lista
        parent=styles['Normal'],
        fontSize=12,
        leftIndent=20,
        spaceAfter=3
    ))
    
    # Elementos para el documento
    elements = []
    
    # Procesar el markdown por líneas
    lines = markdown_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Título principal (H1)
        if line.startswith('# '):
            title = line[2:].strip()
            elements.append(Paragraph(title, styles['Title']))
            elements.append(Spacer(1, 12))
        
        # Subtítulo (H2)
        elif line.startswith('## '):
            subtitle = line[3:].strip()
            elements.append(Paragraph(subtitle, styles['Heading2']))
            elements.append(Spacer(1, 8))
        
        # Subtítulo nivel 3 (H3)
        elif line.startswith('### '):
            subsubtitle = line[4:].strip()
            elements.append(Paragraph(f"<b>{subsubtitle}</b>", styles['Normal']))
            elements.append(Spacer(1, 6))
        
        # Listas
        elif line.startswith('* ') or line.startswith('- '):
            item_text = line[2:].strip()
            # Usar el nuevo estilo para listas
            elements.append(Paragraph(f"• {item_text}", styles['BookListItem']))
        
        # Párrafo normal
        elif line:
            # Procesar texto con formato básico (negrita, cursiva)
            processed_line = line
            
            # Negrita: **texto** -> <b>texto</b>
            processed_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', processed_line)
            
            # Cursiva: *texto* -> <i>texto</i>
            processed_line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', processed_line)
            
            elements.append(Paragraph(processed_line, styles['Normal']))
        
        # Línea en blanco
        elif i > 0 and lines[i-1].strip() and i < len(lines)-1 and lines[i+1].strip():
            elements.append(Spacer(1, 6))
        
        i += 1
    
    # Construir el documento
    doc.build(elements)
    
    return output_file_path

def convert_markdown_to_docx(markdown_content: str, book: Book) -> str:
    """
    Convierte contenido markdown a DOCX usando python-docx.
    
    Args:
        markdown_content: Contenido en formato markdown.
        book: Objeto Book con los datos del libro.
        
    Returns:
        Ruta del archivo DOCX generado.
    """
    
    output_file_path = os.path.join(settings.BOOKS_DIR, f"{book.id}.docx")
    
    # Crear el documento
    doc = Document()
    
    # Configurar márgenes del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Procesar el markdown por líneas
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        # Título principal (H1)
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_paragraph(title)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.runs[0]
            run.font.size = Pt(24)
        
        # Subtítulo (H2)
        elif line.startswith('## '):
            subtitle = line[3:].strip()
            p = doc.add_paragraph(subtitle)
            run = p.runs[0]
            run.font.size = Pt(18)
        
        # Subtítulo nivel 3 (H3)
        elif line.startswith('### '):
            subsubtitle = line[4:].strip()
            p = doc.add_paragraph(subsubtitle)
            run = p.runs[0]
            run.font.size = Pt(16)
        
        # Listas
        elif line.startswith('* ') or line.startswith('- '):
            item_text = line[2:].strip()
            doc.add_paragraph(f"• {item_text}")
        
        # Párrafo normal
        elif line:
            processed_line = line
            
            # Negrita: **texto** -> <b>texto</b>
            processed_line = re.sub(r'\*\*(.*?)\*\*', r'\1', processed_line)
            
            # Cursiva: *texto* -> <i>texto</i>
            processed_line = re.sub(r'\*(.*?)\*', r'\1', processed_line)
            
            doc.add_paragraph(processed_line)
    
    # Guardar el documento
    doc.save(output_file_path)
    
    return output_file_path

def convert_markdown_to_download_file(book: Book, format: FormatDownload) -> str:
    """
    Convierte un archivo markdown a un archivo descargable.
    
    Args:
        book: Objeto Book con la información del libro.
        format: Formato de descarga deseado.
        
    Returns:
        Ruta del archivo descargable.
    """
    # Generamos el contenido markdown
    markdown_content = convert_json_to_markdown(book)
    
    if format == FormatDownload.HTML:
        # Creamos el archivo HTML
        html_content = convert_markdown_to_html(markdown_content, book)
        output_file_path = os.path.join(settings.BOOKS_DIR, f"{book.id}.html")
        with open(output_file_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        return output_file_path
    
    elif format == FormatDownload.PDF:
        # Usamos ReportLab para generar el PDF
        return convert_markdown_to_pdf_with_reportlab(markdown_content, book)
    
    elif format == FormatDownload.MARKDOWN:
        # Guardar el contenido markdown en un archivo
        output_file_path = os.path.join(settings.BOOKS_DIR, f"{book.id}.md")
        with open(output_file_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        return output_file_path
    
    elif format == FormatDownload.TXT:
        # Guardar el contenido markdown en un archivo de texto plano
        output_file_path = os.path.join(settings.BOOKS_DIR, f"{book.id}.txt")
        
        # Simplificar el markdown para texto plano
        plain_text = markdown_content
        plain_text = re.sub(r'# (.*?)$', r'\1\n', plain_text, flags=re.MULTILINE)  # Título
        plain_text = re.sub(r'## (.*?)$', r'\1\n', plain_text, flags=re.MULTILINE)  # Subtítulo
        plain_text = re.sub(r'### (.*?)$', r'\1\n', plain_text, flags=re.MULTILINE)  # Sub-subtítulo
        plain_text = re.sub(r'\*\*(.*?)\*\*', r'\1', plain_text)  # Negrita
        plain_text = re.sub(r'\*(.*?)\*', r'\1', plain_text)  # Cursiva
        plain_text = re.sub(r'- (.*?)$', r'  • \1', plain_text, flags=re.MULTILINE)  # Listas
        
        with open(output_file_path, "w", encoding="utf-8") as f:
            f.write(plain_text)
        return output_file_path
    
    elif format == FormatDownload.DOCX:
        # Usar la función específica para DOCX
        return convert_markdown_to_docx(markdown_content, book)
    
    else:
        raise ValueError(f"Formato no soportado: {format}")